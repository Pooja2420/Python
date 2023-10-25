# Importing document module to save dates in docx extension files.
from docx import Document
# Creating a function to check date is a palindrome or not.
def is_palindrome(s):
    # If the date is equal to date when read from left to right and right to left are same then returns true else returns false.
    return s == s[::-1]

# Creating main function to start the execution of program.
def pal_date(start,end):
# Creating empty list to store the palindrome dates.
    dates = []
    # A loop to iterate over the years from 2001 to 2100 where 2001 is inclsive and 2100 is not inclusive so the loop stop when the range becomes 2099.
    for year in range(start,end+1):
        # A loop to iterate over the months from 1 to 13 where 1 is inclsive and 13 is not inclusive so the loop stop when the range becomes 12.
        for month in range(1, 13):
            # Creating variable to store the number of days in each month
            # Assuming that February has 28 days in every year
            if month in [1, 3, 5, 7, 8, 10, 12] :
                days = 31 
            elif month in [4, 6, 9, 11]:
                days=30
            else:
                days= 28
            # Iterating from from 1 to days + 1 since the last number will not ne inclusive in range function.
            for day in range(1, days + 1):
                # A variable to store the date as a string in DD/MM/YYYY format
                true_date=(f"{day:02d}/{month:02d}/{year:04d}")
                # A variable to store the date as a string in DDMMYYYY format
                date = f"{day:02d}{month:02d}{year:04d}"
               #Calling the is_palindrome function.
                if is_palindrome(date):
                    # If the date is a palindrome, append it to the list. 
                    dates.append(true_date)
    return dates

# Creating the save_dae function to store the palindrame dates in the document.
def save_date(dates):
    # Creating a new Document object
    document = Document()
    # Appending heading of the file.
    document.add_heading("List of Palindrome Dates In 21TH Century", level=2)
    # Looping to add palindrome dates to the document. 
    for date in dates:
        # Adding palindrome dates to the document. 
        document.add_paragraph(date)
    # Save the file with extension.
    document.save("List_of_pal_dates.docx")

    # Printing the total number of palindromes found from 2000 to 2099
    print(f"The program has found {len(dates)} palindrome dates and saved them in List_of_pal_dates.docx.")

# Creating main function to start execution of the program.
def main():
    # Calling the pal_date functions with some arguments to take the dates in proper alignment.
    Start_exec = pal_date(2000, 2099)
    # Calling the save_date function to save the palindrome dates in file.
    save_date(Start_exec)
# Calling main function to start the execution of the program.
main()

